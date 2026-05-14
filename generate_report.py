"""
generate_report.py
==================
Đặt file này vào THƯ MỤC GỐC của dự án.
Cursor sẽ điền dữ liệu vào biến REPORT_DATA bên dưới rồi chạy:
    python generate_report.py
File OUTPUT sẽ được lưu tại: PROJECT_REPORT.docx
"""

import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════
#  DỮ LIỆU — CURSOR SẼ TỰ ĐIỀN VÀO ĐÂY SAU KHI PHÂN TÍCH XONG
# ══════════════════════════════════════════════════════════════
REPORT_DATA = {

    # ── Thông tin dự án ──────────────────────────────────────
    "project_name": "TÊN_DỰ_ÁN",          # Cursor điền
    "project_desc": "Mô tả ngắn dự án",    # Cursor điền
    "analyzed_at":  "",                     # Cursor điền, VD: "14/05/2025 09:30"

    # ── Công nghệ phát hiện được ─────────────────────────────
    # Mỗi phần tử: { "nhom": "...", "ten": "...", "phienban": "...", "ghichu": "..." }
    "technologies": [
        # Cursor tự điền sau khi quét package.json / requirements.txt / go.mod...
    ],

    # ── Luồng hoạt động chatbot ──────────────────────────────
    # Mỗi phần tử: { "buoc": 1, "ten": "...", "mo_ta": "...", "file": "..." }
    "flow_steps": [
        # Cursor điền theo thứ tự thực tế trong code
    ],

    # ── Lỗi & vấn đề phát hiện ──────────────────────────────
    # Loại: "BUG" | "LOGIC" | "DEAD_CODE" | "SECURITY" | "PERFORMANCE"
    # Mức: "CRITICAL" | "HIGH" | "MEDIUM" | "LOW"
    "issues": [
        # { "loai": "BUG", "muc": "HIGH", "file": "src/api/chat.ts",
        #   "dong": 42, "mo_ta": "...", "goi_y_sua": "..." }
    ],

    # ── Tính năng đã hoàn thành ──────────────────────────────
    "completed": [
        # "Giao diện chat cơ bản",
        # "Kết nối OpenAI API",
    ],

    # ── Việc cần làm tiếp theo ───────────────────────────────
    # Mỗi phần tử: { "uu_tien": "P1/P2/P3", "viec": "...", "ly_do": "..." }
    "todo": [
        # { "uu_tien": "P1", "viec": "Thêm rate limiting", "ly_do": "Tránh lạm dụng API" }
    ],

    # ── Đánh giá tổng thể ────────────────────────────────────
    "score":    0,     # 1-10, Cursor điền
    "nhan_xet": "",    # Nhận xét tổng thể, Cursor điền
}

# ══════════════════════════════════════════════════════════════
#  PHẦN CÒN LẠI: KHÔNG CẦN CHỈNH SỬA
# ══════════════════════════════════════════════════════════════

COLORS = {
    "primary":   RGBColor(0x1F, 0x49, 0x7D),
    "secondary": RGBColor(0x2E, 0x74, 0xB5),
    "success":   RGBColor(0x37, 0x86, 0x44),
    "warning":   RGBColor(0xD6, 0x8A, 0x00),
    "danger":    RGBColor(0xC0, 0x39, 0x2B),
    "gray":      RGBColor(0x70, 0x70, 0x70),
    "light":     "EBF3FB",
    "header_bg": "D5E8F0",
    "red_bg":    "FDECEA",
    "green_bg":  "E9F7EF",
    "yellow_bg": "FEF9E7",
}

ISSUE_LEVEL_COLOR = {
    "CRITICAL": ("C0392B", "FDECEA"),
    "HIGH":     ("E67E22", "FEF5E7"),
    "MEDIUM":   ("D6A000", "FEF9E7"),
    "LOW":      ("27AE60", "E9F7EF"),
}
ISSUE_TYPE_ICON = {
    "BUG":         "🐛",
    "LOGIC":       "⚠️",
    "DEAD_CODE":   "🗑️",
    "SECURITY":    "🔒",
    "PERFORMANCE": "⚡",
}
PRIORITY_COLOR = {
    "P1": ("C0392B", "FDECEA"),
    "P2": ("E67E22", "FEF5E7"),
    "P3": ("27AE60", "E9F7EF"),
}


def shade_cell(cell, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_col_widths(table, widths_dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"),    str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")


def add_border_bottom(paragraph, color="1F497D", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Document helpers ──────────────────────────────────────────

def make_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = 11906   # A4
    sec.page_height   = 16838
    sec.left_margin   = Inches(1)
    sec.right_margin  = Inches(1)
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(0.8)
    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    return doc


def h1(doc, text):
    p = doc.add_paragraph()
    add_border_bottom(p, "1F497D", 8)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = COLORS["primary"]
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = COLORS["secondary"]
    return p


def body(doc, text, bold=False, italic=False, color=None, indent_inches=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    if indent_inches:
        p.paragraph_format.left_indent = Inches(indent_inches)
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(10.5)
    if color:
        r.font.color.rgb = color
    return p


def spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)


# ── Table builder ─────────────────────────────────────────────

def make_table(doc, headers, rows, col_widths, header_bg="D5E8F0",
               row_colors=None, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_col_widths(table, col_widths)

    # Header row
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_bg)
        cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(hdr)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = COLORS["primary"]

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = (row_colors[ri] if row_colors and ri < len(row_colors) else None)
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            if bg:
                shade_cell(cell, bg)
            cell_margins(cell)
            p = cell.paragraphs[0]
            r = p.add_run(str(cell_text))
            r.font.size = Pt(font_size)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# ══════════════════════════════════════════════════════════════
#  BUILD REPORT
# ══════════════════════════════════════════════════════════════

def build_report(data):
    doc = make_doc()
    now = data.get("analyzed_at") or datetime.datetime.now().strftime("%d/%m/%Y %H:%M")

    # ── COVER ────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("📋  BÁO CÁO TỔNG HỢP DỰ ÁN")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = COLORS["primary"]

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(data.get("project_name", ""))
    rs.font.size = Pt(15)
    rs.font.color.rgb = COLORS["secondary"]
    rs.bold = True

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = desc.add_run(data.get("project_desc", ""))
    rd.italic = True
    rd.font.size = Pt(11)
    rd.font.color.rgb = COLORS["gray"]

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Phân tích lúc: {now}").font.size = Pt(10)

    spacer(doc, 10)

    # Score banner
    score = data.get("score", 0)
    score_color = "27AE60" if score >= 7 else ("E67E22" if score >= 5 else "C0392B")
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = banner.add_run(f"  Điểm chất lượng code:  {score}/10  ")
    rb.bold = True
    rb.font.size = Pt(16)
    rb.font.color.rgb = RGBColor.from_string(score_color)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    banner._p.get_or_add_pPr().append(shd)
    banner.paragraph_format.space_before = Pt(6)
    banner.paragraph_format.space_after  = Pt(4)

    nhan_xet = data.get("nhan_xet", "")
    if nhan_xet:
        nx = doc.add_paragraph()
        nx.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rnx = nx.add_run(nhan_xet)
        rnx.italic = True
        rnx.font.size = Pt(10.5)
        rnx.font.color.rgb = COLORS["gray"]

    # ── 1. TỔNG QUAN ─────────────────────────────────────────
    h1(doc, "1.  TỔNG QUAN PHÂN TÍCH")

    issues = data.get("issues", [])
    completed = data.get("completed", [])
    todo = data.get("todo", [])
    techs = data.get("technologies", [])

    counts = {"BUG": 0, "LOGIC": 0, "DEAD_CODE": 0, "SECURITY": 0, "PERFORMANCE": 0}
    for iss in issues:
        t = iss.get("loai", "BUG")
        counts[t] = counts.get(t, 0) + 1

    make_table(doc,
        headers=["Hạng mục", "Số lượng", "Ghi chú"],
        rows=[
            ["🐛 Lỗi Bug (crash/runtime)", str(counts["BUG"]),         "Cần sửa ngay"],
            ["⚠️  Lỗi Logic",              str(counts["LOGIC"]),        "Ảnh hưởng chức năng"],
            ["🗑️  Code thừa / Dead code",  str(counts["DEAD_CODE"]),    "Nên dọn dẹp"],
            ["🔒 Vấn đề bảo mật",          str(counts["SECURITY"]),     "Ưu tiên cao"],
            ["⚡ Performance",              str(counts["PERFORMANCE"]),  "Tối ưu thêm"],
            ["✅ Tính năng hoàn thành",     str(len(completed)),         "Đang hoạt động"],
            ["📋 Việc cần làm tiếp",        str(len(todo)),              "Xem mục 4"],
            ["🔧 Công nghệ phát hiện",      str(len(techs)),             "Xem mục 5"],
        ],
        col_widths=[4000, 1800, 3560],
    )

    # ── 2. TÍNH NĂNG ĐÃ HOÀN THÀNH ───────────────────────────
    h1(doc, "2.  TÍNH NĂNG ĐÃ HOÀN THÀNH  ✅")
    if completed:
        make_table(doc,
            headers=["#", "Tính năng / Việc đã làm"],
            rows=[[str(i+1), item] for i, item in enumerate(completed)],
            col_widths=[600, 8760],
            row_colors=["E9F7EF" if i % 2 == 0 else "FFFFFF" for i in range(len(completed))],
        )
    else:
        body(doc, "Chưa có dữ liệu — Cursor sẽ điền sau khi phân tích.", italic=True, color=COLORS["gray"])

    # ── 3. LỖI & VẤN ĐỀ ─────────────────────────────────────
    h1(doc, "3.  LỖI & VẤN ĐỀ PHÁT HIỆN  🔍")

    if issues:
        # Group by severity
        for level in ["CRITICAL", "HIGH", "MEDIUM", "LOW"]:
            group = [x for x in issues if x.get("muc") == level]
            if not group:
                continue
            txt_color, _ = ISSUE_LEVEL_COLOR[level]
            h2(doc, f"  {level}  ({len(group)} vấn đề)")
            rows = []
            bgs  = []
            for iss in group:
                icon = ISSUE_TYPE_ICON.get(iss.get("loai","BUG"), "⚠️")
                rows.append([
                    f"{icon} {iss.get('loai','')}",
                    iss.get("file", ""),
                    f"Dòng {iss.get('dong','')}",
                    iss.get("mo_ta", ""),
                    iss.get("goi_y_sua", ""),
                ])
                _, bg = ISSUE_LEVEL_COLOR[level]
                bgs.append(bg)
            make_table(doc,
                headers=["Loại", "File", "Dòng", "Mô tả lỗi", "Gợi ý sửa"],
                rows=rows,
                col_widths=[1200, 2000, 700, 2800, 2660],
                row_colors=bgs,
                font_size=9.5,
            )
    else:
        body(doc, "Chưa có dữ liệu lỗi — Cursor sẽ điền sau khi phân tích.", italic=True, color=COLORS["gray"])

    # ── 4. VIỆC CẦN LÀM TIẾP THEO ────────────────────────────
    h1(doc, "4.  VIỆC CẦN LÀM TIẾP THEO  📋")

    if todo:
        rows, bgs = [], []
        for i, item in enumerate(todo):
            p = item.get("uu_tien", "P3")
            _, bg = PRIORITY_COLOR.get(p, ("888888", "F8F8F8"))
            rows.append([p, item.get("viec",""), item.get("ly_do","")])
            bgs.append(bg)
        make_table(doc,
            headers=["Ưu tiên", "Công việc cần làm", "Lý do / Ghi chú"],
            rows=rows,
            col_widths=[900, 5000, 3460],
            row_colors=bgs,
        )
        body(doc, "P1 = Làm ngay  |  P2 = Sprint tiếp  |  P3 = Backlog",
             italic=True, color=COLORS["gray"])
    else:
        body(doc, "Chưa có dữ liệu — Cursor sẽ điền sau khi phân tích.", italic=True, color=COLORS["gray"])

    # ── 5. CÔNG NGHỆ SỬ DỤNG ─────────────────────────────────
    h1(doc, "5.  CÔNG NGHỆ SỬ DỤNG  🔧")

    if techs:
        make_table(doc,
            headers=["Nhóm", "Tên", "Phiên bản", "Ghi chú"],
            rows=[[t.get("nhom",""), t.get("ten",""), t.get("phienban",""), t.get("ghichu","")] for t in techs],
            col_widths=[1800, 2400, 1400, 3760],
            row_colors=["EBF3FB" if i % 2 == 0 else "FFFFFF" for i in range(len(techs))],
        )
    else:
        body(doc, "Chưa có dữ liệu — Cursor sẽ điền sau khi phân tích.", italic=True, color=COLORS["gray"])

    # ── 6. LUỒNG HOẠT ĐỘNG CHATBOT ───────────────────────────
    h1(doc, "6.  LUỒNG HOẠT ĐỘNG CHATBOT  🔄")

    flow = data.get("flow_steps", [])
    if flow:
        rows = []
        for step in flow:
            rows.append([
                str(step.get("buoc", "")),
                step.get("ten", ""),
                step.get("mo_ta", ""),
                step.get("file", ""),
            ])
        make_table(doc,
            headers=["#", "Bước", "Mô tả chi tiết", "File / Component"],
            rows=rows,
            col_widths=[500, 2000, 4500, 2360],
            row_colors=["EBF3FB" if i % 2 == 0 else "FFFFFF" for i in range(len(rows))],
        )
    else:
        body(doc, "Chưa có dữ liệu — Cursor sẽ điền sau khi phân tích.", italic=True, color=COLORS["gray"])

    # ── FOOTER ───────────────────────────────────────────────
    spacer(doc, 12)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_bottom(fp, "CCCCCC", 4)
    rfp = fp.add_run(f"Tự động sinh bởi generate_report.py  ·  {now}  ·  Chỉnh sửa thủ công nếu cần")
    rfp.italic = True
    rfp.font.size = Pt(9)
    rfp.font.color.rgb = COLORS["gray"]

    return doc


# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    doc = build_report(REPORT_DATA)
    out = "PROJECT_REPORT.docx"
    doc.save(out)
    print(f"✅ Báo cáo đã được lưu: {out}")
